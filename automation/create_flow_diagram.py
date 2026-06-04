#!/usr/bin/env python3
"""Generate Lead Lifecycle Flow Diagram"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# ============ TITLE ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Duet with Claude — Lead Lifecycle & Call Flow')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Steve Maxim | AAA Life Specialist | Branch Z012')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# ============ FLOW DIAGRAM AS TABLE ============
# Stage colors
stages = [
    {'name': 'NEW', 'color': '3498DB', 'text_color': 'FFFFFF',
     'time': 'Calls 1-3', 'interval': 'Same day\n(AM → PM → EVE)',
     'goal': 'Make first contact',
     'actions': 'No Answer / Left VM\n→ stays here',
     'advance': 'Connected →\nDiscovery →\nAppt Set →',
     'exit': 'After 3 calls\n→ Attempting'},
    {'name': 'ATTEMPTING', 'color': '9B59B6', 'text_color': 'FFFFFF',
     'time': 'Calls 4-10', 'interval': '2-day rest\nbetween cycles\n(AM→PM→EVE)',
     'goal': 'Keep trying to reach',
     'actions': 'No Answer / Left VM\n→ stays here',
     'advance': 'Connected →\nDiscovery →\nAppt Set →',
     'exit': 'After 10 calls\nwith no contact\n→ Auto-Lost'},
    {'name': 'CONNECTED', 'color': '2980B9', 'text_color': 'FFFFFF',
     'time': '5 follow-ups', 'interval': 'Every 2 days',
     'goal': 'Qualify & do\nneeds analysis',
     'actions': 'No Answer / Left VM\n→ stays here',
     'advance': 'Discovery →\nAppt Set →',
     'exit': '5 follow-ups\nno progress\n→ Auto-Nurturing'},
    {'name': 'DISCOVERY', 'color': '1D9E75', 'text_color': 'FFFFFF',
     'time': '5 follow-ups', 'interval': 'Every 3 days',
     'goal': 'Schedule a\npresentation meeting',
     'actions': 'No Answer / Left VM\n→ stays here',
     'advance': 'Appt Set →',
     'exit': '5 follow-ups\nno progress\n→ Auto-Nurturing'},
    {'name': 'MEETING\nSET', 'color': '534AB7', 'text_color': 'FFFFFF',
     'time': '6 follow-ups', 'interval': 'Every 3 days\n+ auto-remind\n1 day before',
     'goal': 'Confirm, present,\nget to a quote',
     'actions': 'Confirmed → stays\nRescheduled → new date\nNo Answer / Left VM',
     'advance': 'Quoted →',
     'exit': '6 follow-ups\nno progress\n→ Auto-Nurturing'},
    {'name': 'QUOTED', 'color': '8E44AD', 'text_color': 'FFFFFF',
     'time': '6 follow-ups', 'interval': 'Every 2 days',
     'goal': 'Get the app signed',
     'actions': 'No Answer / Left VM\n→ stays here\nAppt Set → re-meet',
     'advance': 'App Submitted →',
     'exit': '6 follow-ups\nno progress\n→ Auto-Nurturing'},
    {'name': 'APP\nSUBMITTED', 'color': '27AE60', 'text_color': 'FFFFFF',
     'time': 'Until issued', 'interval': 'As needed',
     'goal': 'Policy issues\nand gets paid',
     'actions': 'No Answer / Left VM\n→ stays here',
     'advance': 'Won →',
     'exit': 'Declined\n→ Lost'},
    {'name': 'WON', 'color': '1D9E75', 'text_color': 'FFFFFF',
     'time': 'Done', 'interval': '—',
     'goal': 'Deal closed!\nIn Book of Business',
     'actions': '—',
     'advance': '—',
     'exit': '—'},
]

# Create the main flow table
table = doc.add_table(rows=1, cols=8)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
widths = [Inches(1.1), Inches(1.0), Inches(1.1), Inches(1.2), Inches(1.6), Inches(1.6), Inches(1.1), Inches(1.1)]

# Header row
headers = ['Stage', 'Timeframe', 'Call\nInterval', 'Goal', 'Dispositions\nthat keep you here', 'Dispositions\nthat advance', 'Auto-Exit\nif no progress', '']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers[:7]):
    hdr_cells[i].text = ''
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(h)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = hdr_cells[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): '333333'
    })
    shading.append(shading_elm)

# Remove last column header
table.rows[0].cells[7].merge(table.rows[0].cells[6])

# Data rows
for s in stages:
    row = table.add_row()
    cells = row.cells

    # Stage name with color
    cells[0].text = ''
    p = cells[0].paragraphs[0]
    run = p.add_run(s['name'])
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = cells[0]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): s['color']
    })
    shading.append(shading_elm)

    # Other cells
    for idx, key in enumerate(['time', 'interval', 'goal', 'actions', 'advance', 'exit']):
        cells[idx+1].text = ''
        p = cells[idx+1].paragraphs[0]
        run = p.add_run(s[key])
        run.font.size = Pt(8)
        if key == 'advance':
            run.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)
            run.font.bold = True
        elif key == 'exit':
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Merge last two columns for exit
    cells[7].merge(cells[6])

# Add cell padding
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

doc.add_paragraph('')

# ============ NURTURING BOX ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NURTURING — The Safety Net')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0xEF, 0x9F, 0x27)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Leads land here when they hit max follow-ups OR you manually move them. They resurface in your call queue based on temperature:')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Nurturing temperature table
nt = doc.add_table(rows=1, cols=4)
nt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = nt.rows[0].cells
for i, h in enumerate(['Temperature', 'Resurfaces Every', 'Use When', 'Can Re-enter']):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    run = p.add_run(h)
    run.font.size = Pt(9)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = hdr[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'EF9F27'
    })
    shading.append(shading_elm)

temps = [
    ('HOT', 'Every 3 days', 'Almost ready, just needs a nudge', 'Connected, Discovery, Appt Set'),
    ('WARM', 'Every 7 days', 'Interested but timing not right', 'Connected, Discovery, Appt Set'),
    ('COOL', 'Every 14 days', 'Long-term prospect, stay top of mind', 'Connected, Discovery, Appt Set'),
]
for name, freq, use, reenter in temps:
    row = nt.add_row().cells
    for i, val in enumerate([name, freq, use, reenter]):
        row[i].text = ''
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 0:
            run.font.bold = True
            color = 'C0392B' if name == 'HOT' else 'EF9F27' if name == 'WARM' else '9B59B6'
            run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ============ DAILY CALL PRIORITY ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Daily Briefing — Call Priority Order')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Your call queue is automatically sorted — closest to money first:')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

pt = doc.add_table(rows=1, cols=3)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = pt.rows[0].cells
for i, h in enumerate(['Priority', 'Stage', 'Why First']):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    run = p.add_run(h)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = hdr[i]._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '534AB7'
    })
    shading.append(shading_elm)

priorities = [
    ('1', 'Quoted (due)', 'Closest to a signature — follow up NOW'),
    ('2', 'Meeting Set (due)', 'Confirm or present — money is near'),
    ('3', 'New Leads', 'Fresh leads — strike while hot'),
    ('4', 'Attempting (due)', 'Keep the cadence going'),
    ('5', 'Connected (due)', 'They know you — advance to discovery'),
    ('6', 'Discovery (due)', 'Qualified — get the meeting booked'),
    ('7', 'Nurturing - Hot', 'Almost ready to re-engage (every 3 days)'),
    ('8', 'Nurturing - Warm', 'Check in (every 7 days)'),
    ('9', 'Nurturing - Cool', 'Stay on radar (every 14 days)'),
]
for pri, stage, why in priorities:
    row = pt.add_row().cells
    for i, val in enumerate([pri, stage, why]):
        row[i].text = ''
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 0:
            run.font.bold = True
            run.font.size = Pt(12)
        if i == 1:
            run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph('')

# ============ FOOTER ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Goal: 40 dials/day | Nurturing = Lost → Nurturing = Future Revenue')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Duet with Claude v1.3.0 — Generated March 30, 2026')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

output_path = '/Users/steve/Desktop/Duet/Lead_Lifecycle_Flow.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
